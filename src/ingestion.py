import os
import re
from pathlib import Path
from dataclasses import dataclass
from typing import List, Dict, Any

@dataclass
class Document:
    text: str
    source: str
    page: int = 1
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        self.metadata["source"] = self.source
        self.metadata["page"] = self.page

def clean_text(text: str) -> str:
    """Strip basic headers, footers, page numbers, and excess whitespace."""
    # Remove multiple newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Remove common page number formats (e.g., Page 1, - 1 -, [1]) at start/end of lines
    text = re.sub(r'(?m)^\s*(Page\s+\d+|\d+\s*of\s*\d+|-\s*\d+\s*-|\[\d+\])\s*$', '', text)
    # Remove leading/trailing whitespace
    return text.strip()

def load_pdf(file_path: Path) -> List[Document]:
    """Load text from a PDF file page by page using pypdf."""
    import pypdf
    documents = []
    try:
        reader = pypdf.PdfReader(file_path)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text() or ""
            cleaned_text = clean_text(text)
            if cleaned_text:
                documents.append(
                    Document(
                        text=cleaned_text,
                        source=file_path.name,
                        page=page_num + 1
                    )
                )
    except Exception as e:
        print(f"Error loading PDF {file_path}: {e}")
    return documents

def load_docx(file_path: Path) -> List[Document]:
    """Load text from a DOCX file using python-docx."""
    import docx
    documents = []
    try:
        doc = docx.Document(file_path)
        # For word documents, we'll treat the entire doc or logical paragraphs as chunks.
        # Since word documents don't have strict page numbers in the text extraction easily,
        # we will group paragraphs into logical "pages" or just treat the whole document as page 1.
        paragraphs_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs_text.append(p.text)
        
        full_text = "\n\n".join(paragraphs_text)
        cleaned_text = clean_text(full_text)
        if cleaned_text:
            documents.append(
                Document(
                    text=cleaned_text,
                    source=file_path.name,
                    page=1
                )
            )
    except Exception as e:
        print(f"Error loading DOCX {file_path}: {e}")
    return documents

def load_txt(file_path: Path) -> List[Document]:
    """Load text from a TXT file."""
    documents = []
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        cleaned_text = clean_text(text)
        if cleaned_text:
            documents.append(
                Document(
                    text=cleaned_text,
                    source=file_path.name,
                    page=1
                )
            )
    except Exception as e:
        print(f"Error loading TXT {file_path}: {e}")
    return documents

def load_documents_from_folder(folder_path: Path) -> List[Document]:
    """Load all supported documents from a folder."""
    all_docs = []
    if not folder_path.exists():
        print(f"Data folder {folder_path} does not exist.")
        return all_docs

    for file_path in folder_path.iterdir():
        if file_path.is_file():
            suffix = file_path.suffix.lower()
            if suffix == ".pdf":
                print(f"Loading PDF: {file_path.name}")
                all_docs.extend(load_pdf(file_path))
            elif suffix == ".docx":
                print(f"Loading DOCX: {file_path.name}")
                all_docs.extend(load_docx(file_path))
            elif suffix in [".txt", ".md"]:
                print(f"Loading Text: {file_path.name}")
                all_docs.extend(load_txt(file_path))
            else:
                print(f"Skipping unsupported file: {file_path.name}")
    return all_docs
