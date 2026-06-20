import os
from pathlib import Path

# Create data directory
data_dir = Path("data")
data_dir.mkdir(exist_ok=True)

# 1. Generate text files first
doc_space = """# Milestones in Space Exploration: From Sputnik to Mars

Space exploration represents one of the greatest technological achievements of humanity. Over the past seventy years, we have transitioned from sending simple radio signals from orbit to planning manned missions to Mars and deploying deep-space telescopes that peer back to the dawn of the universe.

## The Early Space Age (1957–1969)
The modern space race began on October 4, 1957, when the Soviet Union launched Sputnik 1, the first artificial satellite, into orbit. This event shocked the Western world and catalyzed the establishment of NASA in 1958. 

The Soviet Union achieved several more "firsts," including sending Yuri Gagarin into space aboard Vostok 1 in 1961. However, the United States focused its efforts on the Apollo program, aimed at landing a human on the Moon. This culminated on July 20, 1969, when Apollo 11 astronauts Neil Armstrong and Buzz Aldrin stepped onto the lunar surface, declaring it "one giant leap for mankind."

## The Shuttle Era and the ISS (1981–2011)
Following the Apollo missions, NASA shifted focus to reusable spacecraft. The Space Shuttle program, launched in 1981, enabled the deployment of the Hubble Space Telescope in 1990 and the construction of the International Space Space (ISS) beginning in 1998. The ISS, a collaborative effort between NASA, Roscosmos, ESA, JAXA, and CSA, has been continuously occupied since November 2000, serving as a critical microgravity laboratory.

## The Commercial Space Revolution (2010s–Present)
In the 21st century, space exploration transitioned from purely government-funded initiatives to a dynamic commercial sector. Companies like SpaceX, Blue Origin, and Rocket Lab introduced reusable rockets, dramatically lowering the cost of reaching orbit. SpaceX's Falcon 9 became the workhorse of global satellite delivery, and their Dragon spacecraft restored NASA's capability to launch astronauts from American soil.

## Future Horizons: Artemis and Mars
Today, NASA's Artemis program aims to return humans to the Moon, including the first woman and person of color, using the Space Launch System (SLS) and the Orion spacecraft. Artemis is designed to establish a sustainable presence on the Moon and serve as a stepping stone for the ultimate goal: sending humans to Mars by the late 2030s. Robotic missions, such as the Perseverance rover on Mars and the James Webb Space Telescope, continue to expand our scientific understanding of the cosmos daily.
"""

doc_ethics = """# Artificial Intelligence Ethics: Bias, Safety, and Regulation

As Artificial Intelligence (AI) systems become deeply integrated into society, their ethical implications have moved from theoretical concerns to urgent challenges. AI systems influence lending decisions, medical diagnoses, criminal sentencing, and job hiring, making it critical that these models operate fairly, transparently, and safely.

## Algorithmic Bias and Fairness
AI models learn from historical data. If the training data contains human biases, the AI is likely to replicate or even amplify those biases. For instance, natural language processing models have been shown to associate certain occupations with specific genders, while computer vision systems have exhibited lower accuracy rates when identifying individuals with darker skin tones. Addressing algorithmic bias requires diverse datasets, rigorous pre-training audits, and the implementation of mathematical fairness metrics during model optimization.

## Transparency and Explainability (XAI)
Deep learning neural networks are often described as "black boxes" because their decision-making processes are highly complex and difficult for humans to interpret. In critical sectors like healthcare and finance, explainability is essential. Explainable AI (XAI) aims to build models that can explain their reasoning in human-understandable terms. This allows doctors to verify a computer-aided diagnosis and users to understand why a loan application was rejected.

## AI Safety and the Alignment Problem
AI safety focuses on ensuring that advanced AI systems act in accordance with human values and intentions—a challenge known as the alignment problem. If an AI system becomes highly capable but pursues a goal that is misaligned with human well-being, the consequences could be severe. Safety researchers focus on creating robust evaluation frameworks, fail-safe mechanisms, and reinforcement learning systems that incorporate human feedback (RLHF).

## Policy and Global Regulations
Governments worldwide are scrambling to regulate AI. The European Union's AI Act represents the first comprehensive framework, categorizing AI applications by risk levels (from minimal risk to unacceptable risk) and placing strict compliance demands on high-risk use cases like biometric identification. In the United States, executive orders and congressional hearings highlight a growing bipartisan interest in establishing safety standards, protecting user privacy, and preventing AI-driven disinformation.
"""

doc_climate = """# Climate Change Mitigation: Strategies for a Sustainable Future

Climate change is one of the defining crises of our era. To prevent the most catastrophic impacts of global warming, global greenhouse gas emissions must be reduced to net-zero by the mid-21st century. Achieving this goal requires a massive transformation of energy, transportation, agriculture, and industrial systems.

## Transition to Renewable Energy
The combustion of fossil fuels for electricity and heating is the largest source of global emissions. Transitioning to renewable energy sources—such as solar, wind, hydroelectric, and geothermal energy—is the cornerstone of mitigation efforts. Due to rapid technological advances, utility-scale solar and wind energy are now the cheapest sources of new electricity generation in most parts of the world. Grid modernization and battery storage technologies are critical to managing the intermittency of these clean energy sources.

## Decarbonizing Transportation
The transportation sector is heavily dependent on petroleum. The transition to electric vehicles (EVs) is accelerating globally, supported by government mandates, improvements in battery range, and expanding charging infrastructure. For harder-to-decarbonize sectors like aviation, shipping, and heavy trucking, researchers are exploring sustainable aviation fuels (SAFs), green hydrogen, and advanced biofuels.

## Sustainable Agriculture and Land Use
Agriculture, forestry, and land-use changes account for nearly a quarter of global emissions. Key strategies include reducing deforestation, restoring degraded lands, and adopting regenerative farming practices that sequester carbon in the soil. Additionally, reducing food waste and shifting toward plant-based diets can significantly lower emissions associated with livestock production.

## Carbon Capture and Carbon Dioxide Removal (CDR)
While emission reductions are paramount, scientific models suggest that active carbon removal will be necessary to stay within safe temperature limits. Direct Air Capture (DAC) and bioenergy with carbon capture and storage (BECCS) are technological approaches designed to extract CO2 directly from the atmosphere and store it permanently underground. Nature-based solutions, such as reforestation and ocean alkalinization, also play a vital role.
"""

with open(data_dir / "space_exploration_milestones.txt", "w", encoding="utf-8") as f:
    f.write(doc_space)

with open(data_dir / "artificial_intelligence_ethics.txt", "w", encoding="utf-8") as f:
    f.write(doc_ethics)

with open(data_dir / "climate_change_mitigation_strategies.txt", "w", encoding="utf-8") as f:
    f.write(doc_climate)

print("Text files generated.")

# 2. Generate DOCX file
import docx
doc = docx.Document()
doc.add_heading("Python Programming Fundamentals: A Comprehensive Guide", level=1)

p1 = doc.add_paragraph(
    "Python is an interpreted, high-level, general-purpose programming language. Created by Guido van Rossum and first "
    "released in 1991, Python's design philosophy emphasizes code readability with its notable use of significant whitespace. "
    "Its language constructs and object-oriented approach aim to help programmers write clear, logical code for small "
    "and large-scale projects."
)

doc.add_heading("1. Variables and Dynamic Typing", level=2)
p2 = doc.add_paragraph(
    "Unlike statically typed languages like Java or C++, Python uses dynamic typing. This means you do not need to explicitly "
    "declare the variable type before using it; Python infers the type at runtime based on the value assigned. For example, "
    "writing `x = 10` makes `x` an integer, and writing `x = 'Hello'` reassigns it to a string. While this speeds up development, "
    "it requires developers to write rigorous unit tests to avoid runtime type errors."
)

doc.add_heading("2. Built-in Data Structures", level=2)
p3 = doc.add_paragraph(
    "Python includes several powerful built-in data structures:\n"
    "• Lists: Ordered, mutable collections of items, written with square brackets (e.g., [1, 2, 3]).\n"
    "• Dictionaries: Unordered key-value stores, written with curly braces (e.g., {'name': 'Alice', 'age': 25}). Key lookups are O(1) on average.\n"
    "• Sets: Unordered collections of unique elements, useful for membership testing and mathematical operations like union and intersection.\n"
    "• Tuples: Ordered, immutable sequences of elements, written with parentheses (e.g., (1, 2, 3))."
)

doc.add_heading("3. Object-Oriented Programming (OOP)", level=2)
p4 = doc.add_paragraph(
    "Python fully supports Object-Oriented Programming. Classes are defined using the `class` keyword. The constructor is "
    "defined as `__init__`, which accepts `self` as the first argument representing the instance. Python supports multiple inheritance, "
    "operator overloading, and polymorphism. Under the hood, Python objects are backed by dictionaries, which store their attributes."
)

doc.add_heading("4. Advanced Constructs: Decorators and Generators", level=2)
p5 = doc.add_paragraph(
    "Advanced Python features allow developers to write elegant and performant code. Decorators are functions that modify the "
    "behavior of another function or class. They use the `@decorator_name` syntax. Generators are functions that return an "
    "iterator using the `yield` keyword. Unlike standard functions that return a full collection at once, generators yield "
    "values one at a time, making them highly memory-efficient when working with massive streams of data."
)

doc.save(data_dir / "python_programming_fundamentals.docx")
print("DOCX file generated.")

# 3. Generate PDF file using ReportLab
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    pdf_path = data_dir / "transformer_architecture_guide.pdf"
    doc_pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        spaceAfter=15
    )
    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['BodyText'],
        fontSize=10,
        leading=14,
        spaceAfter=10
    )

    story.append(Paragraph("The Transformer Architecture: Attention Is All You Need", title_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph(
        "In 2017, Vaswani et al. published the seminal paper 'Attention Is All You Need', introducing the Transformer architecture. "
        "Before the Transformer, sequence transduction models relied heavily on complex recurrent neural networks (RNNs) or convolutional "
        "neural networks (CNNs) in encoder-decoder structures. These models faced significant limits: recurrent networks process sequences "
        "sequentially, which prevents parallelization during training and makes it difficult to model long-range dependencies.",
        body_style
    ))

    story.append(Paragraph("1. The Core Innovation: Self-Attention Mechanism", h2_style))
    story.append(Paragraph(
        "The fundamental innovation of the Transformer is the self-attention mechanism. Unlike recurrent networks that pass "
        "a hidden state step-by-step, self-attention allows the model to relate different positions of a single sequence "
        "directly. For each word (token) in the input, the model computes three vectors: Query (Q), Key (K), and Value (V). "
        "The attention weights are calculated using the Scaled Dot-Product Attention formula:\n\n"
        "Attention(Q, K, V) = softmax( (Q * K^T) / sqrt(d_k) ) * V\n\n"
        "where d_k is the dimension of the keys. This scaling factor prevents the dot products from growing excessively large, "
        "which would push the softmax function into regions with extremely small gradients.",
        body_style
    ))

    story.append(Paragraph("2. Multi-Head Attention", h2_style))
    story.append(Paragraph(
        "Instead of performing a single attention function, the Transformer utilizes Multi-Head Attention. This involves "
        "linearly projecting the Queries, Keys, and Values h times with different, learned projections. The attention function "
        "is run in parallel on each projection, and the resulting outputs are concatenated and projected again. This allows the "
        "model to jointly attend to information from different representation subspaces at different positions.",
        body_style
    ))

    story.append(Paragraph("3. Encoder-Decoder Structure", h2_style))
    story.append(Paragraph(
        "The encoder consists of a stack of N=6 identical layers. Each layer has two sub-layers: a multi-head self-attention mechanism "
        "and a simple, position-wise fully connected feed-forward network. Residual connections are applied around each sub-layer, "
        "followed by layer normalization. The decoder also contains a stack of N=6 identical layers, but inserts a third sub-layer "
        "which performs multi-head attention over the output of the encoder stack. The self-attention sub-layers in the decoder are "
        "masked to prevent positions from attending to subsequent positions (preserving auto-regressive properties).",
        body_style
    ))

    story.append(Paragraph("4. Positional Encodings", h2_style))
    story.append(Paragraph(
        "Since the Transformer contains no recurrence or convolution, it has no inherent sense of the order or position of tokens "
        "in the sequence. To inject positional information, the authors add 'positional encodings' to the input embeddings. "
        "These encodings utilize sine and cosine functions of different frequencies, allowing the model to easily learn to attend "
        "by relative positions.",
        body_style
    ))

    doc_pdf.build(story)
    print("PDF file generated.")
except Exception as e:
    print(f"Failed to generate PDF using reportlab: {e}")
    print("Please ensure reportlab is installed.")

